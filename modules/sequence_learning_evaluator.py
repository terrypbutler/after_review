"""Evaluate curriculum and pupil access across a sequence of learning."""

from io import BytesIO
import hashlib
import json

import docx
from PIL import Image
import pypdf
import streamlit as st

from config import ANALYSIS_MODEL
from modules.app_secrets import get_secret
from modules import gemini_client as genai
from modules.data_utils import get_ai_response_profile


MAX_SOURCE_CHARS = 80000
MAX_FILE_CHARS = 20000
MAX_IMAGES = 8


def _clean_text(value):
    return " ".join(str(value or "").split())


def _source_fingerprint(uploaded_files, sequence_outline, expected_lessons):
    digest = hashlib.sha256()
    digest.update(str(sequence_outline or "").encode("utf-8"))
    digest.update(str(expected_lessons or 0).encode("utf-8"))
    for uploaded_file in uploaded_files or []:
        digest.update(str(uploaded_file.name).encode("utf-8"))
        digest.update(uploaded_file.getvalue())
    return digest.hexdigest()


def _extract_docx_text(data):
    document = docx.Document(BytesIO(data))
    sections = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            sections.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(section for section in sections if section.strip())


def _extract_pdf_text(data):
    reader = pypdf.PdfReader(BytesIO(data))
    return "\n".join(
        page.extract_text() or ""
        for page in reader.pages
    )


def extract_sequence_sources(uploaded_files, sequence_outline=""):
    """Extract ordered text and images from a sequence or several lesson plans."""
    text_sections = []
    images = []
    notes = []

    if str(sequence_outline or "").strip():
        text_sections.append(
            "TEACHER-PROVIDED SEQUENCE ORDER / CONTEXT:\n"
            + str(sequence_outline).strip()
        )

    for index, uploaded_file in enumerate(uploaded_files or [], start=1):
        filename = str(uploaded_file.name)
        extension = filename.rsplit(".", 1)[-1].casefold()
        data = uploaded_file.getvalue()
        heading = f"SOURCE {index} — {filename}"

        try:
            if extension in {"png", "jpg", "jpeg"}:
                if len(images) >= MAX_IMAGES:
                    notes.append(
                        f"{filename}: skipped because the image limit is {MAX_IMAGES}."
                    )
                    continue
                with Image.open(BytesIO(data)) as source_image:
                    images.append(source_image.copy())
                text_sections.append(
                    f"{heading}\nThis source is supplied as an image in sequence position {index}."
                )
                notes.append(f"{filename}: image ready for analysis.")
                continue

            if extension == "pdf":
                extracted = _extract_pdf_text(data)
            elif extension == "docx":
                extracted = _extract_docx_text(data)
            elif extension == "txt":
                extracted = data.decode("utf-8", errors="replace")
            else:
                notes.append(f"{filename}: unsupported file type.")
                continue

            extracted = extracted.strip()
            if not extracted:
                notes.append(
                    f"{filename}: no readable text was found; a scanned PDF may need "
                    "uploading as images."
                )
                continue
            if len(extracted) > MAX_FILE_CHARS:
                extracted = extracted[:MAX_FILE_CHARS]
                notes.append(
                    f"{filename}: text was shortened to {MAX_FILE_CHARS:,} characters."
                )
            else:
                notes.append(
                    f"{filename}: extracted {len(extracted):,} characters."
                )
            text_sections.append(f"{heading}\n{extracted}")
        except Exception as exc:
            notes.append(f"{filename}: could not be read ({exc}).")

    combined_text = "\n\n".join(text_sections)
    if len(combined_text) > MAX_SOURCE_CHARS:
        combined_text = combined_text[:MAX_SOURCE_CHARS]
        notes.append(
            f"Combined text was shortened to {MAX_SOURCE_CHARS:,} characters."
        )
    return combined_text, images, notes


def build_sequence_prompt(
    source_text,
    profiles_text,
    cohort,
    subject,
    expected_lessons=0,
):
    """Build the structured sequence-evaluation prompt."""
    if cohort == "Year 7":
        age_context = "11 to 12 years old"
        key_stage = "Key Stage 3 (KS3)"
    else:
        age_context = "14 to 15 years old"
        key_stage = "Key Stage 4 (KS4 / GCSE)"

    lesson_hint = (
        f"The teacher expects approximately {expected_lessons} lessons."
        if expected_lessons
        else "Infer the number and order of lessons from the supplied material."
    )

    return f"""
    **[FICTIONAL TEACHER-TRAINING SCENARIO — PUPIL DATA IS MOCK/SYNTHETIC]**
    You are an expert UK curriculum mentor analysing a sequence of learning from
    outside the learner rather than pretending to be a pupil. Your output identifies
    plausible risks and response predictions for rehearsal; it does not prove what any
    pupil knows or whether the sequence will work.

    Subject: {subject}
    Cohort: {cohort} ({age_context})
    Curriculum stage: {key_stage}
    {lesson_hint}

    ORDERED SEQUENCE MATERIAL:
    {source_text or "The sequence is supplied in the attached images, in upload order."}

    COMPACT, PRIVACY-MINIMISED PUPIL RESPONSE PROFILES:
    {profiles_text}

    EPISTEMIC RULES:
    - Every named-pupil response is a profile-informed prediction, not observed evidence
      and not a guarantee about what the pupil will say.
    - "Probable" means strongly supported by the sequence demands and pupil profile.
      "Possible" means a credible alternative worth checking or rehearsing.
    - Predictions must be age-authentic and may be secure, partial, misconception, slip,
      uncertain or no attempt. Include "I don't know", fragments or silence where the
      profile and task demands make that plausible.
    - Do not make pupils uniformly articulate, compliant or successful.

    EVALUATION REQUIREMENTS:
    1. Curriculum intent and endpoint: identify what pupils should know or be able to
       do by the end, and judge alignment with the {key_stage} {subject} curriculum.
    2. Progression and prerequisites: test whether each lesson activates prerequisite
       knowledge, introduces a manageable step and prepares pupils for what follows.
    3. Cumulative learning: evaluate retrieval, spaced practice, deliberate practice,
       vocabulary development, misconception checks and opportunities to apply learning.
    4. Assessment: check whether formative assessment reveals what was learned before
       the next lesson depends on it, and whether the final assessment matches the intent.
    5. Cognitive load: identify abrupt jumps, overloaded lessons, duplication, missing
       bridges and places where scaffolds should fade.
    6. Class access: use the supplied profiles to predict plausible response forms at
       precise points. Identify probable and possible misconceptions, including when a
       pupil may offer a partial response or "I don't know". Preserve high expectations
       and recommend scaffolding, not separate low-ceiling tasks.
    7. Be evidence-led. If a lesson, assessment or curriculum element is absent from the
       supplied material, say it is not visible rather than inventing it.
    8. For each prediction, state what real-pupil evidence would confirm or disconfirm it.

    ANTI-PATTERN GUARDRAILS:
    - Do not use VAK learning styles, left/right-brain claims or the Learning Pyramid.
    - Do not recommend "mild, spicy, hot" tasks or fixed ability worksheets.
    - Do not infer needs from a protected characteristic or SEND label alone.
    - Do not expose private profile text or safeguarding/home-life information.

    Return ONLY valid JSON with this exact top-level structure:
    {{
      "metrics": {{
        "lessons_detected": <integer>,
        "coherence_score": <0-100 integer>,
        "progression_score": <0-100 integer>,
        "retrieval_score": <0-100 integer>,
        "assessment_score": <0-100 integer>,
        "class_access_score": <0-100 integer>,
        "verdict": "<1-3 words>"
      }},
      "overview": "<2-3 sentences explaining the sequence's overall learning journey>",
      "curriculum": {{
        "intended_endpoint": "<what pupils should know/do by the end>",
        "coverage": "<coverage and pitch judgement>",
        "missing_or_unclear": "<important curriculum content not visible>"
      }},
      "sequence_map": [
        {{
          "lesson": <integer>,
          "title": "<short inferred or supplied title>",
          "new_learning": "<main new knowledge or skill>",
          "builds_on": "<prerequisite or previous learning>",
          "retrieval": "<what is retrieved, or 'Not visible'>",
          "assessment": "<how understanding is checked, or 'Not visible'>",
          "bridge_forward": "<how this prepares for the next lesson>",
          "risk": "<main sequencing or access risk>"
        }}
      ],
      "gaps_and_bottlenecks": [
        {{
          "location": "<lesson number or across sequence>",
          "gap": "<specific learning gap>",
          "impact": "<why this matters later>",
          "repair": "<precise change>"
        }}
      ],
      "response_predictions": [
        {{
          "name": "<exact pupil name from the supplied profiles>",
          "profile_basis": "<non-sensitive qualities relevant to this prediction>",
          "most_likely_response": "<age-authentic predicted response form>",
          "probable_misconception": "<specific misconception or 'None strongly indicated'>",
          "possible_misconception": "<credible alternative misconception>",
          "likelihood": "<Probable or Possible>",
          "confidence": "<Low, Medium or High>",
          "risk_point": "<specific lesson/transition>",
          "adjustment": "<high-expectation scaffold, probe or extension>",
          "verify_live": "<question, work sample or observation needed from real pupils>"
        }}
      ],
      "priority_actions": [
        {{
          "priority": "<High, Medium or Low>",
          "location": "<where to change the sequence>",
          "action": "<specific actionable revision>",
          "reason": "<learning benefit>"
        }}
      ],
      "validation_plan": [
        {{
          "location": "<lesson/task>",
          "check": "<hinge question, exit ticket or observable evidence>",
          "prediction_tested": "<which risk or misconception this checks>"
        }}
      ]
    }}

    Include every detected lesson in sequence_map. Select 4-6 distinct pupils for
    response_predictions, representing contrasting attainment, confidence,
    participation, processing and access patterns rather than merely selecting pupils
    with SEND.
    """


def _bounded_score(value):
    try:
        return max(0, min(100, int(round(float(value)))))
    except (TypeError, ValueError):
        return 0


def normalise_sequence_result(result, allowed_names):
    """Constrain model output before rendering it in the app."""
    clean_result = result if isinstance(result, dict) else {}
    metrics = clean_result.get("metrics", {})
    if not isinstance(metrics, dict):
        metrics = {}
    clean_result["metrics"] = {
        "lessons_detected": max(
            0,
            int(_bounded_score(metrics.get("lessons_detected", 0))),
        ),
        "coherence_score": _bounded_score(metrics.get("coherence_score", 0)),
        "progression_score": _bounded_score(metrics.get("progression_score", 0)),
        "retrieval_score": _bounded_score(metrics.get("retrieval_score", 0)),
        "assessment_score": _bounded_score(metrics.get("assessment_score", 0)),
        "class_access_score": _bounded_score(metrics.get("class_access_score", 0)),
        "verdict": _clean_text(metrics.get("verdict", "Review needed"))[:40],
    }

    for list_key in (
        "sequence_map",
        "gaps_and_bottlenecks",
        "response_predictions",
        "priority_actions",
        "validation_plan",
    ):
        if not isinstance(clean_result.get(list_key), list):
            clean_result[list_key] = []
        clean_result[list_key] = [
            item
            for item in clean_result[list_key]
            if isinstance(item, dict)
        ]

    allowed = {str(name).strip() for name in allowed_names}
    clean_result["response_predictions"] = [
        prediction
        for prediction in clean_result["response_predictions"]
        if str(prediction.get("name", "")).strip() in allowed
    ][:6]
    return clean_result


def _render_score(label, value):
    st.metric(label, f"{value}/100")
    st.progress(value / 100)


def render_sequence_result(result):
    metrics = result.get("metrics", {})
    st.success(
        f"Sequence evaluated: {metrics.get('lessons_detected', 0)} lessons detected."
    )
    st.caption(
        "Scores are heuristic planning indicators. Named responses are "
        "profile-informed predictions for rehearsal, not observed pupil evidence."
    )

    metric_columns = st.columns(6)
    metric_columns[0].metric(
        "Lessons",
        metrics.get("lessons_detected", 0),
    )
    metric_columns[1].metric(
        "Verdict",
        metrics.get("verdict", "Review"),
    )
    metric_columns[2].metric(
        "Coherence",
        f"{metrics.get('coherence_score', 0)}/100",
    )
    metric_columns[3].metric(
        "Progression",
        f"{metrics.get('progression_score', 0)}/100",
    )
    metric_columns[4].metric(
        "Retrieval",
        f"{metrics.get('retrieval_score', 0)}/100",
    )
    metric_columns[5].metric(
        "Assessment",
        f"{metrics.get('assessment_score', 0)}/100",
    )

    st.info(result.get("overview", "No overview was returned."))

    curriculum = result.get("curriculum", {})
    if not isinstance(curriculum, dict):
        curriculum = {}
    st.markdown("### Curriculum intent and endpoint")
    curriculum_columns = st.columns(3)
    curriculum_columns[0].markdown(
        f"**Intended endpoint**\n\n{curriculum.get('intended_endpoint', 'Not visible')}"
    )
    curriculum_columns[1].markdown(
        f"**Coverage and pitch**\n\n{curriculum.get('coverage', 'Not visible')}"
    )
    curriculum_columns[2].markdown(
        f"**Missing or unclear**\n\n{curriculum.get('missing_or_unclear', 'Nothing identified')}"
    )

    st.markdown("### Lesson-to-lesson learning map")
    sequence_map = result.get("sequence_map", [])
    if sequence_map:
        st.dataframe(
            sequence_map,
            width="stretch",
            hide_index=True,
            column_order=[
                "lesson",
                "title",
                "new_learning",
                "builds_on",
                "retrieval",
                "assessment",
                "bridge_forward",
                "risk",
            ],
        )
    else:
        st.warning("No lesson map was returned.")

    gap_column, prediction_column = st.columns(2)
    with gap_column:
        st.markdown("### Gaps and bottlenecks")
        gaps = result.get("gaps_and_bottlenecks", [])
        if not gaps:
            st.caption("No specific gap was returned.")
        for index, gap in enumerate(gaps, start=1):
            with st.expander(
                f"{index}. {gap.get('location', 'Across sequence')} — "
                f"{gap.get('gap', 'Gap')}",
                expanded=index <= 2,
            ):
                st.write(gap.get("impact", ""))
                st.success(gap.get("repair", ""))

    with prediction_column:
        st.markdown("### Profile-informed response predictions")
        predictions = result.get("response_predictions", [])
        if not predictions:
            st.caption("No valid pupil prediction was returned.")
        for prediction in predictions:
            with st.expander(
                f"{prediction.get('name', 'Pupil')} — "
                f"{prediction.get('likelihood', 'Possible')} "
                f"({prediction.get('confidence', 'Medium')} confidence)",
                expanded=False,
            ):
                st.caption(
                    f"Profile basis: {prediction.get('profile_basis', 'Not stated')}"
                )
                st.write(
                    "**Most likely response:** "
                    f"{prediction.get('most_likely_response', 'Not predicted')}"
                )
                st.warning(
                    "**Probable misconception:** "
                    f"{prediction.get('probable_misconception', 'None strongly indicated')}"
                )
                st.info(
                    "**Possible misconception:** "
                    f"{prediction.get('possible_misconception', 'Not identified')}"
                )
                st.caption(
                    f"Risk point: {prediction.get('risk_point', 'Not identified')}"
                )
                st.success(
                    f"Adjustment: {prediction.get('adjustment', 'Not identified')}"
                )
                st.caption(
                    "Verify with real pupils: "
                    f"{prediction.get('verify_live', 'Collect live evidence')}"
                )

    st.markdown("### Priority revisions")
    priorities = result.get("priority_actions", [])
    if not priorities:
        st.caption("No priority revision was returned.")
    for action in priorities:
        priority = str(action.get("priority", "Medium"))
        st.markdown(
            f"**{priority} · {action.get('location', 'Across sequence')}**  \n"
            f"{action.get('action', '')}  \n"
            f"*Why: {action.get('reason', '')}*"
        )

    st.markdown("### What to verify with real pupils")
    validation_plan = result.get("validation_plan", [])
    if not validation_plan:
        st.caption("No validation check was returned.")
    for item in validation_plan:
        st.markdown(
            f"**{item.get('location', 'Lesson')}** — "
            f"{item.get('check', 'Collect pupil evidence')}  \n"
            f"*Tests: {item.get('prediction_tested', 'identified risk')}*"
        )


def render_sequence_evaluator(df, cohort, subject="General"):
    st.subheader(f"🧭 Sequence Evaluator: {cohort} {subject}")
    st.caption(
        "Upload a scheme of work or several lesson plans in teaching order. "
        "The evaluator predicts plausible pressure points and pupil responses for "
        "rehearsal; real-pupil evidence is still required."
    )

    api_key = get_secret("GEMINI_API_KEY")
    if not api_key:
        st.error("⚠️ Gemini API Key missing. Please add it to Streamlit secrets.")
        return
    genai.configure(api_key=api_key)

    st.sidebar.markdown("---")
    st.sidebar.markdown("### 🧭 Sequence evaluation")
    st.sidebar.caption(
        "Reviews curriculum intent, prerequisites, progression, retrieval, "
        "assessment and cognitive load, then proposes profile-informed response "
        "predictions and checks to validate with real pupils."
    )

    uploaded_files = st.file_uploader(
        "Upload the sequence or lesson plans in teaching order",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
        accept_multiple_files=True,
        key="sequence_files",
        help=(
            "You may upload one scheme-of-work file or several lesson plans. "
            "Files are analysed in the order shown."
        ),
    )
    sequence_outline = st.text_area(
        "Sequence order, intended endpoint or pasted outline (optional)",
        key="sequence_outline",
        placeholder=(
            "For example: Lesson 1 establishes..., Lesson 2 applies..., "
            "the final assessment requires..."
        ),
        height=150,
    )
    expected_lessons = st.number_input(
        "Approximate number of lessons (optional)",
        min_value=0,
        max_value=60,
        value=0,
        step=1,
        help="Leave at 0 and the evaluator will infer the number.",
    )

    if uploaded_files:
        st.caption(
            "Teaching order: "
            + " → ".join(uploaded_file.name for uploaded_file in uploaded_files)
        )

    fingerprint = _source_fingerprint(
        uploaded_files,
        sequence_outline,
        expected_lessons,
    )
    stored_context = st.session_state.get("sequence_evaluation_context")
    if stored_context and stored_context != fingerprint:
        st.session_state["sequence_evaluation_result"] = None

    has_source = bool(uploaded_files or sequence_outline.strip())
    if not has_source:
        st.info(
            "Upload at least one document/image or paste a sequence outline to begin."
        )

    if st.button(
        "🚀 Evaluate sequence of learning",
        type="primary",
        width="stretch",
        disabled=not has_source,
    ):
        source_text, image_parts, extraction_notes = extract_sequence_sources(
            uploaded_files,
            sequence_outline,
        )
        with st.expander("Source extraction details", expanded=False):
            for note in extraction_notes:
                st.write(f"- {note}")

        if not source_text and not image_parts:
            st.error("No readable sequence material was found.")
        else:
            profiles = []
            for _, row in df.iterrows():
                name = str(row.get("Full Name", "Unknown"))
                profile = get_ai_response_profile(
                    row,
                    cohort,
                    subject,
                    max_chars=700,
                )
                profiles.append(f"- {name}: {profile}")
            profiles_text = "\n".join(profiles)

            prompt = build_sequence_prompt(
                source_text,
                profiles_text,
                cohort,
                subject,
                expected_lessons,
            )
            with st.spinner(
                "Tracing curriculum progression and profile-informed response "
                "predictions across the sequence..."
            ):
                try:
                    model = genai.GenerativeModel(ANALYSIS_MODEL)
                    contents = [prompt, *image_parts]
                    response = model.generate_content(
                        contents,
                        generation_config={
                            "response_mime_type": "application/json"
                        },
                    )
                    raw_text = (
                        response.text
                        .replace("```json", "")
                        .replace("```", "")
                        .strip()
                    )
                    result = json.loads(raw_text)
                    result = normalise_sequence_result(
                        result,
                        df["Full Name"].astype(str).tolist(),
                    )
                    st.session_state["sequence_evaluation_result"] = result
                    st.session_state["sequence_evaluation_context"] = fingerprint
                except Exception as exc:
                    st.error(f"Could not evaluate the sequence: {exc}")

    result = st.session_state.get("sequence_evaluation_result")
    if (
        result
        and st.session_state.get("sequence_evaluation_context") == fingerprint
    ):
        st.markdown("---")
        render_sequence_result(result)
        st.download_button(
            "⬇️ Download evaluation as JSON",
            data=json.dumps(result, indent=2, ensure_ascii=False),
            file_name=f"{cohort}_{subject}_sequence_evaluation.json".replace(
                " ",
                "_",
            ),
            mime="application/json",
            width="stretch",
        )
